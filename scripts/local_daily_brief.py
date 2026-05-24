#!/usr/bin/env python3
"""
============================================================
  Financial News Collector — Local Daily Brief
============================================================
Runs on Mac at 08:20 via launchd.

Pipeline:
  1. Google News RSS search — 40 authoritative sources
  2. SiliconFlow LLM — 利好/利空 impact analysis
  3. PDF report (ReportLab)
  4. DOCX report (python-docx)
  5. Douyin carousel images (1080×1920)
  6. Save to ~/Desktop/每日财经新闻/YYYY-MM-DD/

Requires: pip3 install requests reportlab python-docx Pillow feedparser
Env: SF_API_KEY (optional, for impact analysis)
"""

import os, sys, time, json, hashlib, re, textwrap, tempfile
from datetime import datetime, timezone, timedelta
from urllib.parse import quote
import urllib.parse
import xml.etree.ElementTree as ET

import requests

BEIJING_TZ = timezone(timedelta(hours=8))
DESKTOP = os.path.expanduser("~/Desktop")

# ============================================================
# 40 Sources
# ============================================================
SOURCES = [
    # Chinese Media (10)
    {"id":"cls","name":"财联社","query":"财联社 最新财经快讯","cat":"breaking_news","lang":"zh","rel":"high"},
    {"id":"sina_finance","name":"新浪财经","query":"新浪财经 最新财经新闻","cat":"comprehensive","lang":"zh","rel":"high"},
    {"id":"eastmoney","name":"东方财富","query":"东方财富 最新财经要闻","cat":"comprehensive","lang":"zh","rel":"high"},
    {"id":"wallstreetcn","name":"华尔街见闻","query":"华尔街见闻 最新资讯","cat":"global_finance","lang":"zh","rel":"high"},
    {"id":"stcn","name":"证券时报","query":"证券时报 最新证券新闻","cat":"securities","lang":"zh","rel":"high"},
    {"id":"cs_com_cn","name":"中国证券报","query":"中国证券报 最新证券资讯","cat":"securities","lang":"zh","rel":"high"},
    {"id":"yicai","name":"第一财经","query":"第一财经 最新财经新闻","cat":"comprehensive","lang":"zh","rel":"high"},
    {"id":"21jingji","name":"21世纪经济报道","query":"21世纪经济报道 最新财经","cat":"comprehensive","lang":"zh","rel":"high"},
    {"id":"chinafundnews","name":"中国基金报","query":"中国基金报 基金新闻","cat":"comprehensive","lang":"zh","rel":"high"},
    {"id":"ce_cn","name":"经济日报","query":"经济日报 最新经济新闻","cat":"comprehensive","lang":"zh","rel":"highest"},
    # International Research (3)
    {"id":"goldman_sachs","name":"高盛研报","query":"Goldman Sachs China research report","cat":"investment_research","lang":"en","rel":"highest"},
    {"id":"jpmorgan","name":"摩根大通研报","query":"JPMorgan China market insights","cat":"investment_research","lang":"en","rel":"highest"},
    {"id":"morgan_stanley","name":"摩根士丹利研报","query":"Morgan Stanley China equity research","cat":"investment_research","lang":"en","rel":"highest"},
    # Domestic Brokerage (4)
    {"id":"citic_research","name":"中信证券研究","query":"中信证券 最新研报 A股策略","cat":"domestic_research","lang":"zh","rel":"highest"},
    {"id":"cicc_research","name":"中金公司研究","query":"中金公司 最新研报 策略","cat":"domestic_research","lang":"zh","rel":"highest"},
    {"id":"htsc_research","name":"华泰证券研究","query":"华泰证券 最新研报 行业分析","cat":"domestic_research","lang":"zh","rel":"highest"},
    {"id":"gtja_research","name":"国泰君安研究","query":"国泰君安 最新研报 策略","cat":"domestic_research","lang":"zh","rel":"highest"},
    # International Media (5)
    {"id":"reuters","name":"路透社","query":"Reuters China finance markets latest","cat":"global_news","lang":"en","rel":"highest"},
    {"id":"bloomberg","name":"彭博社","query":"Bloomberg China economy market news","cat":"global_finance","lang":"en","rel":"highest"},
    {"id":"cnbc","name":"CNBC","query":"CNBC China Asia markets latest","cat":"global_markets","lang":"en","rel":"high"},
    {"id":"wsj","name":"华尔街日报","query":"WSJ China finance economy news","cat":"global_finance","lang":"en","rel":"highest"},
    {"id":"ft","name":"金融时报","query":"Financial Times China markets economy","cat":"global_finance","lang":"en","rel":"highest"},
    # Chinese Government (10)
    {"id":"gov_cn","name":"中国政府网","query":"中国政府网 最新产业政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"ndrc","name":"国家发改委","query":"国家发改委 最新产业政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"miit","name":"工信部","query":"工信部 最新产业政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"most_cn","name":"科技部","query":"科技部 最新科技政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"mofcom","name":"商务部","query":"商务部 最新贸易政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"mof_cn","name":"财政部","query":"财政部 最新财税政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"pbc","name":"中国人民银行","query":"中国人民银行 货币政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"csrc","name":"证监会","query":"证监会 最新监管政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"nfra","name":"金融监管总局","query":"金融监管总局 最新政策","cat":"cn_government","lang":"zh","rel":"highest"},
    {"id":"stats_cn","name":"国家统计局","query":"国家统计局 最新经济数据","cat":"cn_government","lang":"zh","rel":"highest"},
    # International Government (8)
    {"id":"federal_reserve","name":"美联储","query":"Federal Reserve FOMC statement press release","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"us_treasury","name":"美国财政部","query":"US Treasury sanctions policy latest","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"us_sec","name":"美国SEC","query":"US SEC securities regulation China","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"us_commerce","name":"美国商务部","query":"US Commerce export control entity list","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"ustr","name":"USTR","query":"USTR China tariff trade latest","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"ecb","name":"欧洲央行","query":"ECB monetary policy decision press release","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"eu_commission","name":"欧盟委员会","query":"European Commission trade China anti-subsidy","cat":"intl_government","lang":"en","rel":"highest"},
    {"id":"boj","name":"日本央行","query":"Bank of Japan monetary policy statement","cat":"intl_government","lang":"en","rel":"highest"},
]

CAT_LABELS = {
    "breaking_news":"⚡ 快讯/突发新闻","comprehensive":"📋 综合财经新闻",
    "global_finance":"🌐 全球金融市场","global_news":"📡 国际通讯社",
    "global_markets":"📈 全球市场行情","securities":"📜 证券/监管新闻",
    "investment_research":"🏦 国际投行研报","domestic_research":"🏢 国内券商研报",
    "cn_government":"🇨🇳 中国政府产业政策","intl_government":"🌍 国际政府/央行政策",
}
CAT_ORDER = ["breaking_news","cn_government","intl_government","comprehensive",
             "global_news","global_markets","investment_research","domestic_research",
             "global_finance","securities"]

_session = None
def get_session():
    global _session
    if _session is None:
        _session = requests.Session()
        _session.headers.update({"User-Agent":"Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"})
    return _session

def now_bj(): return datetime.now(BEIJING_TZ)
def now_str(): return now_bj().strftime("%Y-%m-%d %H:%M:%S")
def gen_id(*p): return hashlib.md5(":".join(p).encode()).hexdigest()[:12]

# ============================================================
# Google News RSS Search
# ============================================================
def google_news_search(query, max_results=8, retries=3):
    sess = get_session()
    lang = "zh-CN" if any('\u4e00'<=c<='\u9fff' for c in query) else "en-US"
    gl = "CN" if lang=="zh-CN" else "US"
    ceid = f"{gl}:{lang.replace('-','')}"
    url = f"https://news.google.com/rss/search?q={quote(query)}&hl={lang}&gl={gl}&ceid={ceid}"

    for attempt in range(retries):
        try:
            r = sess.get(url, timeout=20)
            if r.status_code != 200: continue
            root = ET.fromstring(r.text)
            items = root.findall(".//item")
            results = []
            for item in items[:max_results]:
                t = item.find("title")
                l = item.find("link")
                d = item.find("description")
                s = item.find("source")
                title = t.text.strip() if t is not None and t.text else ""
                link = l.text.strip() if l is not None and l.text else ""
                desc = d.text.strip() if d is not None and d.text else ""
                source = s.text.strip() if s is not None and s.text else ""
                # Clean HTML
                title = re.sub(r'<[^>]+>','',title).strip()
                desc = re.sub(r'<[^>]+>','',desc).strip()
                if title:
                    results.append({"title":title,"href":link,"body":desc,"source":source})
            if results: return results
        except: pass
        if attempt < retries-1: time.sleep(2)
    return []

# ============================================================
# LLM Impact Analysis
# ============================================================
def analyze_impact(items):
    """DeepSeek V3: filter A-share relevant, summarize, analyze A/HK/US impact."""
    api_key = os.environ.get("SF_API_KEY","")
    if not api_key:
        print("  ⚠️  SF_API_KEY not set, skipping impact analysis")
        return items

    print(f"  DeepSeek V3 analyzing {len(items)} items...")
    batch_size = 15
    analyzed = []

    for start in range(0, len(items), batch_size):
        batch = items[start:start+batch_size]
        lines = "\n".join(f"[{i}] [{it['source_name']}] {it['title']}" for i,it in enumerate(batch))

        prompt = f"""你是资深A股分析师。从以下新闻中筛选出可能对A股市场产生影响的新闻。
对每条重要新闻：(1)一句话概述内容 (2)分别分析对A股、港股、美股的影响。
无影响或中性的不要输出。

输出JSON数组：
[{{"idx":0,"impact":"利好/利空","summary":"新闻概述",
   "a_share":"对A股影响","hk_stock":"对港股影响(无则填无)",
   "us_stock":"对美股影响(无则填无)",
   "sectors":["板块"],"stocks":["代码 名称"]}}]

新闻：
{lines}

只输出JSON数组。"""

        try:
            resp = requests.post("https://api.siliconflow.cn/v1/chat/completions",
                json={"model":"deepseek-ai/DeepSeek-V3",
                      "messages":[{"role":"user","content":prompt}],
                      "temperature":0.3,"max_tokens":4000},
                headers={"Authorization":f"Bearer {api_key}"}, timeout=180)
            ct = resp.json()["choices"][0]["message"]["content"]
            js = ct.find("["); je = ct.rfind("]")+1
            if js>=0 and je>js:
                analyses = json.loads(ct[js:je])
                kept = set()
                for a in analyses:
                    idx = a.get("idx",-1)
                    if 0<=idx<len(batch) and a.get("impact") in ("利好","利空"):
                        batch[idx]["impact"] = a["impact"]
                        batch[idx]["ai_summary"] = a.get("summary","")
                        batch[idx]["a_share"] = a.get("a_share","")
                        batch[idx]["hk_stock"] = a.get("hk_stock","")
                        batch[idx]["us_stock"] = a.get("us_stock","")
                        batch[idx]["sectors"] = ", ".join(a.get("sectors",[]))
                        batch[idx]["stocks"] = ", ".join(a.get("stocks",[]))
                        kept.add(idx)
                kept_batch = [batch[i] for i in range(len(batch)) if i in kept]
                analyzed.extend(kept_batch)
                print(f"    [{start+1}-{min(start+batch_size,len(items))}] kept {len(kept)}/{len(batch)}")
        except Exception as e:
            print(f"    Batch {start} error: {e}")
        if start+batch_size<len(items): time.sleep(1)

    removed = len(items) - len(analyzed)
    print(f"  -> {len(analyzed)} market-relevant, {removed} noise removed")
    return analyzed

def build_pdf(items, date_str, path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, grey
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fn = "Helvetica"
    for fp in ["/System/Library/Fonts/STHeiti Light.ttc","/System/Library/Fonts/PingFang.ttc"]:
        if os.path.exists(fp):
            try: pdfmetrics.registerFont(TTFont("CJK",fp)); fn="CJK"; break
            except: pass
    if fn=="Helvetica":
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light')); fn="STSong-Light"
        except: pass

    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=2*cm,rightMargin=2*cm,topMargin=1.5*cm,bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    ts = ParagraphStyle('T',parent=styles['Title'],fontName=fn,fontSize=16,leading=22,alignment=TA_CENTER,textColor=HexColor('#1a1a2e'))
    ss = ParagraphStyle('S',parent=styles['Normal'],fontName=fn,fontSize=8,textColor=grey,alignment=TA_CENTER,spaceAfter=12)
    hs = ParagraphStyle('H',parent=styles['Heading2'],fontName=fn,fontSize=11,leading=16,textColor=HexColor('#e94560'),spaceBefore=12,spaceAfter=4)
    bs = ParagraphStyle('B',parent=styles['Normal'],fontName=fn,fontSize=8,leading=12,spaceAfter=2)
    rs = ParagraphStyle('R',parent=styles['Normal'],fontName=fn,fontSize=6,leading=8,textColor=grey)

    story = [Spacer(1,0.5*cm), Paragraph("每日财经新闻聚合报告",ts),
             Paragraph(f"日期: {date_str}　新闻: {len(items)} 条",ss),
             HRFlowable(width="100%",thickness=1,color=HexColor('#e94560')), Spacer(1,8)]
    by_cat = {}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        story.append(Paragraph(CAT_LABELS.get(cat,cat),hs))
        for it in by_cat[cat][:6]:
            t = it["title"].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            imp = it.get("impact","")
            if imp: t+=f' <font color="#22c55e">【{imp}】</font>' if "利好" in imp else (f' <font color="#ef4444">【{imp}】</font>' if "利空" in imp else f' 【{imp}】')
            lnk = f'<font color="blue"><u><a href="{it.get("url","")}">{t}</a></u></font>' if it.get("url") else t
            story.append(Paragraph(lnk,bs))
            src = it["source_name"]
            if it.get("ai_summary"): 
                story.append(Paragraph(f"📝 {it['ai_summary']}",rs))
            impact_parts = []
            if it.get("a_share"): impact_parts.append(f"A股: {it['a_share']}")
            if it.get("hk_stock") and it["hk_stock"]!="无": impact_parts.append(f"港股: {it['hk_stock']}")
            if it.get("us_stock") and it["us_stock"]!="无": impact_parts.append(f"美股: {it['us_stock']}")
            if impact_parts:
                story.append(Paragraph(" | ".join(impact_parts), rs))
            if it.get("stocks"): src+=f" | 📈 {it['stocks']}"
            story.append(Paragraph(src,rs))
            story.append(Spacer(1,4))
        story.append(Spacer(1,4))
    story.append(HRFlowable(width="100%",thickness=0.5,color=grey))
    story.append(Paragraph("自动生成 · Financial News Collector · 仅供参考",rs))
    doc.build(story)

# ============================================================
# DOCX
# ============================================================
def build_docx(items, date_str, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    t = doc.add_heading('每日财经新闻聚合报告',0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f'日期: {date_str}　新闻: {len(items)} 条'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    by_cat = {}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        doc.add_heading(CAT_LABELS.get(cat,cat),1)
        for it in by_cat[cat][:6]:
            imp = it.get("impact","")
            if imp:
                ip = doc.add_paragraph(f'【{imp}】')
                ip.runs[0].font.size=Pt(9)
                ip.runs[0].font.color.rgb = RGBColor(34,197,94) if "利好" in imp else (RGBColor(239,68,68) if "利空" in imp else RGBColor(136,136,136))
            pp = doc.add_paragraph(); pp.paragraph_format.space_after=Pt(2)
            r = pp.add_run(f'• {it["title"]}'); r.bold=True; r.font.size=Pt(9)
            ps = doc.add_paragraph(f'  {it["source_name"]}'); ps.runs[0].font.size=Pt(7); ps.runs[0].font.color.rgb=RGBColor(160,160,160)
            if it.get("reasoning"):
                pr = doc.add_paragraph(f'  💡 {it["reasoning"]}'); pr.runs[0].font.size=Pt(8)
            if it.get("stocks"):
                pk = doc.add_paragraph(f'  📈 {it["stocks"]}'); pk.runs[0].font.size=Pt(8)
    doc.save(path)

# ============================================================
# Douyin Images
# ============================================================
def build_douyin(items, date_str, out_dir):
    from PIL import Image, ImageDraw, ImageFont
    W,H = 1080,1920
    paths = []
    for fp in ["/System/Library/Fonts/STHeiti Light.ttc","/System/Library/Fonts/PingFang.ttc"]:
        if os.path.exists(fp):
            try: tf=ImageFont.truetype(fp,56); bf=ImageFont.truetype(fp,32); sf=ImageFont.truetype(fp,24); break
            except: pass
    else: tf=bf=sf=ImageFont.load_default()
    BG,ACC,GOLD,WH,LG,CD = (18,18,36),(233,69,96),(255,200,60),(255,255,255),(180,180,190),(30,30,55)
    # Cover
    img = Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img)
    d.rectangle([(0,0),(W,8)],fill=ACC)
    d.text((80,200),date_str,fill=LG,font=sf)
    tt="每日财经要闻"; tw=d.textbbox((0,0),tt,font=tf)[2]; d.text(((W-tw)//2,280),tt,fill=WH,font=tf)
    sub=f"共 {len(items)} 条 · {len(set(i['source_id'] for i in items))} 个来源"
    sw=d.textbbox((0,0),sub,font=bf)[2]; d.text(((W-sw)//2,380),sub,fill=GOLD,font=bf)
    d.line([(340,460),(740,460)],fill=ACC,width=3)
    by_cat={}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    y=540
    for cat in CAT_ORDER:
        if cat in by_cat: d.text((120,y),f"{CAT_LABELS.get(cat,cat)}　{len(by_cat[cat])} 条",fill=LG,font=sf); y+=50
    cv=os.path.join(out_dir,"douyin_00_cover.png"); img.save(cv,quality=95); paths.append(cv)
    # Category slides
    si=1
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        ci=by_cat[cat][:5]
        img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img)
        d.rectangle([(0,0),(W,160)],fill=CD)
        d.text((80,40),CAT_LABELS.get(cat,cat),fill=ACC,font=tf)
        d.text((80,110),f"{len(ci)} 条重要新闻",fill=LG,font=sf)
        y=220
        for it in ci:
            ch=280
            if y+ch>H-100: ch=H-100-y
            d.rectangle([(40,y),(W-40,y+ch)],fill=CD,outline=(60,60,80))
            imp=it.get("impact","")
            yo=0
            if imp:
                ic=(34,197,94) if "利好" in imp else ((239,68,68) if "利空" in imp else LG)
                d.text((80,y+20),f"【{imp}】",fill=ic,font=sf); yo=50
            d.text((80,y+20),it["source_name"],fill=GOLD,font=sf)
            d.text((80,y+60+yo),textwrap.fill(it["title"],width=30),fill=WH,font=bf)
            y+=ch+30
        fp=os.path.join(out_dir,f"douyin_{si:02d}_{cat}.png"); img.save(fp,quality=95); paths.append(fp); si+=1
    return paths

# ============================================================
# Main
# ============================================================
def main():
    print(f"[{now_str()}] Local Daily Brief starting...")
    yesterday = (now_bj()-timedelta(days=1)).strftime("%Y-%m-%d")
    today = now_bj().strftime("%Y-%m-%d")
    out_dir = os.path.join(DESKTOP, "每日财经新闻", today)
    os.makedirs(out_dir, exist_ok=True)
    print(f"  Output: {out_dir}")

    # 1. Search
    print(f"\n[1/4] Searching {len(SOURCES)} sources...")
    items = []
    for i, src in enumerate(SOURCES):
        print(f"  [{i+1}/{len(SOURCES)}] {src['name']:<12s} ", end="", flush=True)
        raw = google_news_search(src["query"], max_results=5)
        cnt = 0
        for r in raw:
            t = r.get("title",""); h = r.get("href",""); b = r.get("body","")
            if t:
                items.append({"id":gen_id(src["id"],t,h),"title":t[:200],"url":h,"summary":b[:300],
                              "source_id":src["id"],"source_name":src["name"],
                              "category":src["cat"],"lang":src["lang"],"reliability":src["rel"]})
                cnt += 1
        print(f"{cnt}")
        if i < len(SOURCES)-1: time.sleep(2)
    hits = len(set(i["source_id"] for i in items))
    print(f"  -> {len(items)} items from {hits}/{len(SOURCES)} sources")



    # 2. Impact analysis
    print(f"\n[2/4] Market impact analysis...")
    items = analyze_impact(items)

    # 3. PDF + DOCX
    print(f"\n[3/4] Generating files...")
    pdf_path = os.path.join(out_dir, f"每日财经新闻_{today}.pdf")
    build_pdf(items, today, pdf_path)
    print(f"  PDF: {os.path.getsize(pdf_path)/1024:.1f} KB")
    docx_path = os.path.join(out_dir, f"每日财经新闻_{today}.docx")
    build_docx(items, today, docx_path)
    print(f"  DOCX: {os.path.getsize(docx_path)/1024:.1f} KB")

    # 4. Douyin
    dy_dir = os.path.join(out_dir, "douyin")
    os.makedirs(dy_dir, exist_ok=True)
    dy_paths = build_douyin(items, today, dy_dir)
    print(f"  Douyin: {len(dy_paths)} slides")

    # Summary
    with open(os.path.join(out_dir,"summary.txt"),"w") as f:
        f.write(f"每日财经新闻 {today}\n")
        f.write(f"筛选后 {len(items)} 条重要新闻 (已过滤噪音)\n")
        f.write(f"覆盖 {hits}/{len(SOURCES)} 个来源\n\n")
        # Market briefing generated in PDF/DOCX
        f.write("=== 分类统计 ===\n")
        for cat in CAT_ORDER:
            cnt = len([i for i in items if i["category"]==cat])
            if cnt: f.write(f"{CAT_LABELS.get(cat,cat)}: {cnt} 条\n")

    print(f"\n[{now_str()}] Done! → {out_dir}")

if __name__ == "__main__":
    main()

# ============================================================
# PDF
# ============================================================

# ============================================================
# LLM Filter + Summarize
# ============================================================
def build_pdf(items, date_str, path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, grey
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fn = "Helvetica"
    for fp in ["/System/Library/Fonts/STHeiti Light.ttc","/System/Library/Fonts/PingFang.ttc"]:
        if os.path.exists(fp):
            try: pdfmetrics.registerFont(TTFont("CJK",fp)); fn="CJK"; break
            except: pass
    if fn=="Helvetica":
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light')); fn="STSong-Light"
        except: pass

    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=2*cm,rightMargin=2*cm,topMargin=1.5*cm,bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    ts = ParagraphStyle('T',parent=styles['Title'],fontName=fn,fontSize=16,leading=22,alignment=TA_CENTER,textColor=HexColor('#1a1a2e'))
    ss = ParagraphStyle('S',parent=styles['Normal'],fontName=fn,fontSize=8,textColor=grey,alignment=TA_CENTER,spaceAfter=12)
    hs = ParagraphStyle('H',parent=styles['Heading2'],fontName=fn,fontSize=11,leading=16,textColor=HexColor('#e94560'),spaceBefore=12,spaceAfter=4)
    bs = ParagraphStyle('B',parent=styles['Normal'],fontName=fn,fontSize=8,leading=12,spaceAfter=2)
    rs = ParagraphStyle('R',parent=styles['Normal'],fontName=fn,fontSize=6,leading=8,textColor=grey)

    story = [Spacer(1,0.5*cm), Paragraph("每日财经新闻聚合报告",ts),
             Paragraph(f"日期: {date_str}　新闻: {len(items)} 条",ss),
             HRFlowable(width="100%",thickness=1,color=HexColor('#e94560')), Spacer(1,8)]
    by_cat = {}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        story.append(Paragraph(CAT_LABELS.get(cat,cat),hs))
        for it in by_cat[cat][:6]:
            t = it["title"].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            imp = it.get("impact","")
            if imp: t+=f' <font color="#22c55e">【{imp}】</font>' if "利好" in imp else (f' <font color="#ef4444">【{imp}】</font>' if "利空" in imp else f' 【{imp}】')
            lnk = f'<font color="blue"><u><a href="{it.get("url","")}">{t}</a></u></font>' if it.get("url") else t
            story.append(Paragraph(lnk,bs))
            src = it["source_name"]
            if it.get("ai_summary"): 
                story.append(Paragraph(f"📝 {it['ai_summary']}",rs))
            impact_parts = []
            if it.get("a_share"): impact_parts.append(f"A股: {it['a_share']}")
            if it.get("hk_stock") and it["hk_stock"]!="无": impact_parts.append(f"港股: {it['hk_stock']}")
            if it.get("us_stock") and it["us_stock"]!="无": impact_parts.append(f"美股: {it['us_stock']}")
            if impact_parts:
                story.append(Paragraph(" | ".join(impact_parts), rs))
            if it.get("stocks"): src+=f" | 📈 {it['stocks']}"
            story.append(Paragraph(src,rs))
            story.append(Spacer(1,4))
        story.append(Spacer(1,4))
    story.append(HRFlowable(width="100%",thickness=0.5,color=grey))
    story.append(Paragraph("自动生成 · Financial News Collector · 仅供参考",rs))
    doc.build(story)

# ============================================================
# DOCX
# ============================================================
def build_docx(items, date_str, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    t = doc.add_heading('每日财经新闻聚合报告',0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f'日期: {date_str}　新闻: {len(items)} 条'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    by_cat = {}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        doc.add_heading(CAT_LABELS.get(cat,cat),1)
        for it in by_cat[cat][:6]:
            imp = it.get("impact","")
            if imp:
                ip = doc.add_paragraph(f'【{imp}】')
                ip.runs[0].font.size=Pt(9)
                ip.runs[0].font.color.rgb = RGBColor(34,197,94) if "利好" in imp else (RGBColor(239,68,68) if "利空" in imp else RGBColor(136,136,136))
            pp = doc.add_paragraph(); pp.paragraph_format.space_after=Pt(2)
            r = pp.add_run(f'• {it["title"]}'); r.bold=True; r.font.size=Pt(9)
            ps = doc.add_paragraph(f'  {it["source_name"]}'); ps.runs[0].font.size=Pt(7); ps.runs[0].font.color.rgb=RGBColor(160,160,160)
            if it.get("reasoning"):
                pr = doc.add_paragraph(f'  💡 {it["reasoning"]}'); pr.runs[0].font.size=Pt(8)
            if it.get("stocks"):
                pk = doc.add_paragraph(f'  📈 {it["stocks"]}'); pk.runs[0].font.size=Pt(8)
    doc.save(path)

# ============================================================
# Douyin Images
# ============================================================
def build_douyin(items, date_str, out_dir):
    from PIL import Image, ImageDraw, ImageFont
    W,H = 1080,1920
    paths = []
    for fp in ["/System/Library/Fonts/STHeiti Light.ttc","/System/Library/Fonts/PingFang.ttc"]:
        if os.path.exists(fp):
            try: tf=ImageFont.truetype(fp,56); bf=ImageFont.truetype(fp,32); sf=ImageFont.truetype(fp,24); break
            except: pass
    else: tf=bf=sf=ImageFont.load_default()
    BG,ACC,GOLD,WH,LG,CD = (18,18,36),(233,69,96),(255,200,60),(255,255,255),(180,180,190),(30,30,55)
    # Cover
    img = Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img)
    d.rectangle([(0,0),(W,8)],fill=ACC)
    d.text((80,200),date_str,fill=LG,font=sf)
    tt="每日财经要闻"; tw=d.textbbox((0,0),tt,font=tf)[2]; d.text(((W-tw)//2,280),tt,fill=WH,font=tf)
    sub=f"共 {len(items)} 条 · {len(set(i['source_id'] for i in items))} 个来源"
    sw=d.textbbox((0,0),sub,font=bf)[2]; d.text(((W-sw)//2,380),sub,fill=GOLD,font=bf)
    d.line([(340,460),(740,460)],fill=ACC,width=3)
    by_cat={}; [by_cat.setdefault(it["category"],[]).append(it) for it in items]
    y=540
    for cat in CAT_ORDER:
        if cat in by_cat: d.text((120,y),f"{CAT_LABELS.get(cat,cat)}　{len(by_cat[cat])} 条",fill=LG,font=sf); y+=50
    cv=os.path.join(out_dir,"douyin_00_cover.png"); img.save(cv,quality=95); paths.append(cv)
    # Category slides
    si=1
    for cat in CAT_ORDER:
        if cat not in by_cat: continue
        ci=by_cat[cat][:5]
        img=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(img)
        d.rectangle([(0,0),(W,160)],fill=CD)
        d.text((80,40),CAT_LABELS.get(cat,cat),fill=ACC,font=tf)
        d.text((80,110),f"{len(ci)} 条重要新闻",fill=LG,font=sf)
        y=220
        for it in ci:
            ch=280
            if y+ch>H-100: ch=H-100-y
            d.rectangle([(40,y),(W-40,y+ch)],fill=CD,outline=(60,60,80))
            imp=it.get("impact","")
            yo=0
            if imp:
                ic=(34,197,94) if "利好" in imp else ((239,68,68) if "利空" in imp else LG)
                d.text((80,y+20),f"【{imp}】",fill=ic,font=sf); yo=50
            d.text((80,y+20),it["source_name"],fill=GOLD,font=sf)
            d.text((80,y+60+yo),textwrap.fill(it["title"],width=30),fill=WH,font=bf)
            y+=ch+30
        fp=os.path.join(out_dir,f"douyin_{si:02d}_{cat}.png"); img.save(fp,quality=95); paths.append(fp); si+=1
    return paths

# ============================================================
# Main
# ============================================================
def main():
    print(f"[{now_str()}] Local Daily Brief starting...")
    yesterday = (now_bj()-timedelta(days=1)).strftime("%Y-%m-%d")
    today = now_bj().strftime("%Y-%m-%d")
    out_dir = os.path.join(DESKTOP, "每日财经新闻", today)
    os.makedirs(out_dir, exist_ok=True)
    print(f"  Output: {out_dir}")

    # 1. Search
    print(f"\n[1/4] Searching {len(SOURCES)} sources...")
    items = []
    for i, src in enumerate(SOURCES):
        print(f"  [{i+1}/{len(SOURCES)}] {src['name']:<12s} ", end="", flush=True)
        raw = google_news_search(src["query"], max_results=5)
        cnt = 0
        for r in raw:
            t = r.get("title",""); h = r.get("href",""); b = r.get("body","")
            if t:
                items.append({"id":gen_id(src["id"],t,h),"title":t[:200],"url":h,"summary":b[:300],
                              "source_id":src["id"],"source_name":src["name"],
                              "category":src["cat"],"lang":src["lang"],"reliability":src["rel"]})
                cnt += 1
        print(f"{cnt}")
        if i < len(SOURCES)-1: time.sleep(2)
    hits = len(set(i["source_id"] for i in items))
    print(f"  -> {len(items)} items from {hits}/{len(SOURCES)} sources")



    # 2. Impact analysis
    print(f"\n[2/4] Market impact analysis...")
    items = analyze_impact(items)

    # 3. PDF + DOCX
    print(f"\n[3/4] Generating files...")
    pdf_path = os.path.join(out_dir, f"每日财经新闻_{today}.pdf")
    build_pdf(items, today, pdf_path)
    print(f"  PDF: {os.path.getsize(pdf_path)/1024:.1f} KB")
    docx_path = os.path.join(out_dir, f"每日财经新闻_{today}.docx")
    build_docx(items, today, docx_path)
    print(f"  DOCX: {os.path.getsize(docx_path)/1024:.1f} KB")

    # 4. Douyin
    dy_dir = os.path.join(out_dir, "douyin")
    os.makedirs(dy_dir, exist_ok=True)
    dy_paths = build_douyin(items, today, dy_dir)
    print(f"  Douyin: {len(dy_paths)} slides")

    # Summary
    with open(os.path.join(out_dir,"summary.txt"),"w") as f:
        f.write(f"每日财经新闻 {today}\n")
        f.write(f"筛选后 {len(items)} 条重要新闻 (已过滤噪音)\n")
        f.write(f"覆盖 {hits}/{len(SOURCES)} 个来源\n\n")
        # Market briefing generated in PDF/DOCX
        f.write("=== 分类统计 ===\n")
        for cat in CAT_ORDER:
            cnt = len([i for i in items if i["category"]==cat])
            if cnt: f.write(f"{CAT_LABELS.get(cat,cat)}: {cnt} 条\n")

    print(f"\n[{now_str()}] Done! → {out_dir}")

if __name__ == "__main__":
    main()

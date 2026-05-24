#!/usr/bin/env python3
"""
GitHub Actions Daily Financial Brief
=====================================
1. Searches 37 sources via DuckDuckGo HTML
2. Generates PDF report
3. 3. LLM analyzes capital market impact (利好/利空)
4. Generates PDF + DOCX + Douyin images
5. Emails to phone
"""

import os, sys, time, json, hashlib, re, smtplib, tempfile, textwrap
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime, timezone, timedelta
import requests
from urllib.parse import quote

BEIJING_TZ = timezone(timedelta(hours=8))

SOURCES = [
    {"id": "cls", "name": "财联社", "query": "财联社 最新财经快讯", "cat": "breaking_news", "lang": "zh", "rel": "high"},
    {"id": "sina_finance", "name": "新浪财经", "query": "新浪财经 最新财经新闻", "cat": "comprehensive", "lang": "zh", "rel": "high"},
    {"id": "eastmoney", "name": "东方财富", "query": "东方财富 最新财经要闻", "cat": "comprehensive", "lang": "zh", "rel": "high"},
    {"id": "wallstreetcn", "name": "华尔街见闻", "query": "华尔街见闻 最新资讯", "cat": "global_finance", "lang": "zh", "rel": "high"},
    {"id": "stcn", "name": "证券时报", "query": "证券时报 最新证券新闻", "cat": "securities", "lang": "zh", "rel": "high"},
    {"id": "cs_com_cn", "name": "中国证券报", "query": "中国证券报 最新证券资讯", "cat": "securities", "lang": "zh", "rel": "high"},
    {"id": "yicai", "name": "第一财经", "query": "第一财经 最新财经新闻", "cat": "comprehensive", "lang": "zh", "rel": "high"},
    {"id": "21jingji", "name": "21世纪经济报道", "query": "21世纪经济报道 最新财经", "cat": "comprehensive", "lang": "zh", "rel": "high"},
    {"id": "ce_cn", "name": "经济日报", "query": "经济日报 最新经济新闻", "cat": "comprehensive", "lang": "zh", "rel": "highest"},
    {"id": "chinafundnews", "name": "中国基金报", "query": "中国基金报 基金新闻", "cat": "comprehensive", "lang": "zh", "rel": "high"},
    {"id": "goldman_sachs", "name": "高盛研报", "query": "Goldman Sachs China research report", "cat": "investment_research", "lang": "en", "rel": "highest"},
    {"id": "jpmorgan", "name": "摩根大通研报", "query": "JPMorgan China market insights report", "cat": "investment_research", "lang": "en", "rel": "highest"},
    {"id": "morgan_stanley", "name": "摩根士丹利研报", "query": "Morgan Stanley China equity research", "cat": "investment_research", "lang": "en", "rel": "highest"},
    {"id": "citic_research", "name": "中信证券研究", "query": "中信证券 最新研报 A股策略", "cat": "domestic_research", "lang": "zh", "rel": "highest"},
    {"id": "cicc_research", "name": "中金公司研究", "query": "中金公司 最新研报 策略", "cat": "domestic_research", "lang": "zh", "rel": "highest"},
    {"id": "htsc_research", "name": "华泰证券研究", "query": "华泰证券 最新研报 行业分析", "cat": "domestic_research", "lang": "zh", "rel": "highest"},
    {"id": "gtja_research", "name": "国泰君安研究", "query": "国泰君安 最新研报 策略", "cat": "domestic_research", "lang": "zh", "rel": "highest"},
    {"id": "reuters", "name": "路透社", "query": "Reuters China finance markets latest news", "cat": "global_news", "lang": "en", "rel": "highest"},
    {"id": "bloomberg", "name": "彭博社", "query": "Bloomberg China economy market news", "cat": "global_finance", "lang": "en", "rel": "highest"},
    {"id": "cnbc", "name": "CNBC", "query": "CNBC China Asia markets latest", "cat": "global_markets", "lang": "en", "rel": "high"},
    {"id": "wsj", "name": "华尔街日报", "query": "WSJ China finance economy news", "cat": "global_finance", "lang": "en", "rel": "highest"},
    {"id": "ft", "name": "金融时报", "query": "Financial Times China markets economy", "cat": "global_finance", "lang": "en", "rel": "highest"},
    {"id": "gov_cn", "name": "中国政府网", "query": "中国政府网 最新产业政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "ndrc", "name": "国家发改委", "query": "国家发改委 最新产业政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "miit", "name": "工信部", "query": "工信部 最新产业政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "mofcom", "name": "商务部", "query": "商务部 最新贸易政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "mof_cn", "name": "财政部", "query": "财政部 最新财税政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "pbc", "name": "中国人民银行", "query": "中国人民银行 货币政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "csrc", "name": "证监会", "query": "证监会 最新监管政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "nfra", "name": "金融监管总局", "query": "金融监管总局 最新政策", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "stats_cn", "name": "国家统计局", "query": "国家统计局 最新经济数据", "cat": "cn_government", "lang": "zh", "rel": "highest"},
    {"id": "federal_reserve", "name": "美联储", "query": "Federal Reserve FOMC statement press release", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "us_treasury", "name": "美国财政部", "query": "US Treasury sanctions policy latest", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "us_commerce", "name": "美国商务部", "query": "US Commerce Department export control entity list", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "ustr", "name": "USTR", "query": "USTR China tariff trade latest", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "ecb", "name": "欧洲央行", "query": "ECB monetary policy decision press release", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "eu_commission", "name": "欧盟委员会", "query": "European Commission trade China anti-subsidy", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "boe", "name": "英国央行", "query": "Bank of England monetary policy report", "cat": "intl_government", "lang": "en", "rel": "highest"},
    {"id": "boj", "name": "日本央行", "query": "Bank of Japan monetary policy statement", "cat": "intl_government", "lang": "en", "rel": "highest"},
]

CAT_LABELS = {
    "breaking_news": "⚡ 快讯/突发新闻",
    "comprehensive": "📋 综合财经新闻",
    "global_finance": "🌐 全球金融市场",
    "global_news": "📡 国际通讯社",
    "global_markets": "📈 全球市场行情",
    "securities": "📜 证券/监管新闻",
    "investment_research": "🏦 国际投行研报",
    "domestic_research": "🏢 国内券商研报",
    "cn_government": "🇨🇳 中国政府产业政策",
    "intl_government": "🌍 国际政府/央行政策",
}

SMTP_CFG = {
    "qq.com": ("smtp.qq.com", 465, True),
    "gmail.com": ("smtp.gmail.com", 465, True),
    "163.com": ("smtp.163.com", 465, True),
}

SESSION = None

def get_session():
    global SESSION
    if SESSION is None:
        SESSION = requests.Session()
        SESSION.headers.update({
            "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
        })
    return SESSION


def now_bj():
    return datetime.now(BEIJING_TZ)


def gen_id(*parts):
    return hashlib.md5(":".join(parts).encode()).hexdigest()[:12]


def ddg_search(query, max_results=5):
    """Search via DDG HTML (primary) or DDG Lite (fallback)."""
    sess = get_session()

    # Try DDG HTML first
    try:
        url = f"https://html.duckduckgo.com/html/"
        resp = sess.post(url, data={"q": query}, timeout=20)
        if resp.status_code == 200 and len(resp.text) > 500:
            return _parse_ddg_html(resp.text, max_results)
    except Exception:
        pass

    # Fallback: DDG Lite
    try:
        url = f"https://lite.duckduckgo.com/lite/"
        resp = sess.post(url, data={"q": query}, timeout=20)
        if resp.status_code == 200:
            return _parse_ddg_lite(resp.text, max_results)
    except Exception:
        pass

    return []


def _parse_ddg_html(html, max_results):
    results = []
    link_pat = re.compile(
        r'<a[^>]*class="result__a"[^>]*href="([^"]*)"[^>]*>(.*?)</a>',
        re.DOTALL | re.IGNORECASE
    )
    snippet_pat = re.compile(
        r'<a[^>]*class="result__snippet"[^>]*>(.*?)</a>',
        re.DOTALL | re.IGNORECASE
    )
    links = link_pat.findall(html)
    snippets = snippet_pat.findall(html)
    for i, (href, title_raw) in enumerate(links[:max_results]):
        title = re.sub(r'<[^>]+>', '', title_raw).strip()
        if not title:
            continue
        snippet = ""
        if i < len(snippets):
            snippet = re.sub(r'<[^>]+>', '', snippets[i]).strip()
        results.append({"title": title, "href": href.strip(), "body": snippet})
    return results


def _parse_ddg_lite(html, max_results):
    results = []
    # DDG Lite: each result is in <tr class="result-snippet">
    # Link is <a rel="nofollow" href="...">
    rows = re.findall(r'<tr[^>]*class="result-snippet"[^>]*>(.*?)</tr>', html, re.DOTALL)
    for row in rows[:max_results]:
        link_m = re.search(r'href="([^"]+)"', row)
        title_m = re.search(r'<a[^>]*>(.*?)</a>', row)
        snippet_m = re.search(r'class="result-snippet">(.*?)</td>', row, re.DOTALL)
        if title_m:
            title = re.sub(r'<[^>]+>', '', title_m.group(1)).strip()
            href = link_m.group(1) if link_m else ""
            snippet = re.sub(r'<[^>]+>', '', snippet_m.group(1)).strip() if snippet_m else ""
            results.append({"title": title, "href": href, "body": snippet})
    return results


def search_all(since_date, max_per=5):
    all_items = []
    total = len(SOURCES)
    for i, src in enumerate(SOURCES):
        query = src["query"]
        if since_date and "近" not in query and "最新" in query:
            pass  # query already has 最新

        print(f"  [{i+1}/{total}] {src['name']} ", end="", flush=True)
        try:
            raw = ddg_search(query, max_results=max_per)
        except Exception as e:
            print(f"ERR:{e}")
            raw = []

        count = 0
        for r in raw:
            title = r.get("title", "")
            href = r.get("href", "")
            body = r.get("body", "")
            if title:
                all_items.append({
                    "id": gen_id(src["id"], title, href),
                    "title": title[:200],
                    "url": href,
                    "summary": body[:300],
                    "source_id": src["id"],
                    "source_name": src["name"],
                    "category": src["cat"],
                    "lang": src["lang"],
                    "reliability": src["rel"],
                })
                count += 1
        print(f"{count}")
        if i < total - 1:
            time.sleep(2)
    return all_items


def text_report(items, date_str):
    analyzed = sum(1 for i in items if i.get("impact"))
    lines = [f"📰 每日财经要闻", f"日期: {date_str}",
             f"共搜集 {len(items)} 条新闻 · {analyzed} 条已分析市场影响", ""]
    by_cat = {}
    for item in items:
        by_cat.setdefault(item["category"], []).append(item)
    cat_order = ["breaking_news", "cn_government", "intl_government", "comprehensive",
                 "global_news", "global_markets", "investment_research", "domestic_research"]
    for cat in cat_order:
        if cat in by_cat:
            lines.append(f"\n{CAT_LABELS.get(cat, cat)}")
            for item in by_cat[cat][:5]:
                impact_tag = f" {item.get('impact','')}" if item.get('impact') else ""
                lines.append(f"  • [{item['source_name']}]{impact_tag} {item['title'][:100]}")
                if item.get("reasoning"):
                    lines.append(f"    💡 {item['reasoning'][:120]}")
                if item.get("stocks"):
                    lines.append(f"    📈 关注: {item['stocks'][:150]}")
                if item.get("url"):
                    lines.append(f"    {item['url']}")
    hits = len(set(i["source_id"] for i in items))
    lines.append(f"\n\n📊 覆盖 {hits}/{len(SOURCES)} 个来源")
    lines.append("\n---\n自动生成 · GitHub Actions")
    return "\n".join(lines)


def build_pdf(items, date_str, pdf_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, grey
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "Helvetica"
    for fp in ["/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
               "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"]:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("CJK", fp))
                font_name = "CJK"
                break
            except:
                pass

    if font_name == "Helvetica":
        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            font_name = "STSong-Light"
        except:
            pass

    doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()

    title_s = ParagraphStyle('T', parent=styles['Title'], fontName=font_name,
                              fontSize=16, leading=22, spaceAfter=4, alignment=TA_CENTER,
                              textColor=HexColor('#1a1a2e'))
    sub_s = ParagraphStyle('S', parent=styles['Normal'], fontName=font_name,
                            fontSize=8, leading=10, textColor=grey, alignment=TA_CENTER, spaceAfter=12)
    h2_s = ParagraphStyle('H', parent=styles['Heading2'], fontName=font_name,
                           fontSize=11, leading=16, textColor=HexColor('#e94560'),
                           spaceBefore=12, spaceAfter=4)
    body_s = ParagraphStyle('B', parent=styles['Normal'], fontName=font_name,
                             fontSize=8, leading=12, spaceAfter=2, alignment=TA_JUSTIFY)
    src_s = ParagraphStyle('R', parent=styles['Normal'], fontName=font_name,
                            fontSize=6, leading=8, textColor=grey)

    story = [Spacer(1, 0.5*cm),
             Paragraph("每日财经新闻聚合报告", title_s),
             Paragraph(f"日期: {date_str}　|　新闻: {len(items)} 条", sub_s),
             HRFlowable(width="100%", thickness=1, color=HexColor('#e94560')),
             Spacer(1, 8)]

    by_cat = {}
    for item in items:
        by_cat.setdefault(item["category"], []).append(item)
    cat_order = ["breaking_news", "cn_government", "intl_government", "comprehensive",
                 "global_news", "global_markets", "investment_research", "domestic_research"]

    for cat in cat_order:
        if cat not in by_cat:
            continue
        story.append(Paragraph(CAT_LABELS.get(cat, cat), h2_s))
        for item in by_cat[cat][:6]:
            title = item["title"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            url = item.get("url", "")
            summary = item.get("summary", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            source = item["source_name"]

            impact = item.get("impact", "")
            if impact:
                impact_color = "#22c55e" if "利好" in impact else ("#ef4444" if "利空" in impact else "#888")
                title = f'{title} <font color="{impact_color}">【{impact}】</font>'
            if url:
                link = f'<font color="blue"><u><a href="{url}">{title}</a></u></font>'
            else:
                link = title
            story.append(Paragraph(link, body_s))
            reasoning = item.get("reasoning", "")
            src_line = source
            if reasoning:
                src_line += f" | 💡 {reasoning}"
            story.append(Paragraph(src_line, src_s))
            if summary and summary != title:
                story.append(Paragraph(summary[:200], body_s))
            story.append(Spacer(1, 4))
        story.append(Spacer(1, 4))

    story.append(HRFlowable(width="100%", thickness=0.5, color=grey))
    story.append(Spacer(1, 6))
    story.append(Paragraph("自动生成 · GitHub Actions · 仅供参考", src_s))

    doc.build(story)





def analyze_impact(items):
    """Use SiliconFlow LLM to analyze capital market impact of each news item.
    Returns items enriched with impact, sectors, stocks, reasoning."""
    import urllib.request as urlreq

    api_key = os.environ.get("SF_API_KEY", "")
    if not api_key:
        print("  ⚠️  SF_API_KEY not set, skipping impact analysis")
        return items

    print(f"  Analyzing {len(items)} items with SiliconFlow LLM...")

    # Process in batches to stay within token limits
    batch_size = 20
    analyzed = []

    for batch_start in range(0, len(items), batch_size):
        batch = items[batch_start:batch_start + batch_size]
        batch_text = ""
        for i, item in enumerate(batch):
            batch_text += f"[{i}] {item['source_name']}: {item['title']}\n"

        prompt = f"""Analyze the following Chinese & global financial news headlines.
For EACH item [N], output a JSON array with one object per item:
[
  {{"idx": 0, "impact": "利好/利空/中性", "direction": "positive/negative/neutral",
    "sectors": ["affected sector1", "sector2"],
    "stocks": ["stock1 ticker", "stock2 ticker"],
    "reasoning": "1-sentence market impact in Chinese"}}
]

News:
{batch_text}

Output ONLY valid JSON array, no other text."""

        payload = json.dumps({
            "model": "Qwen/Qwen2.5-7B-Instruct",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 4000,
        }).encode()

        req = urlreq.Request("https://api.siliconflow.cn/openai/v1/chat/completions",
                             data=payload,
                             headers={"Authorization": f"Bearer {api_key}",
                                      "Content-Type": "application/json"})
        try:
            with urlreq.urlopen(req, timeout=60) as resp:
                result = json.loads(resp.read())
            content = result["choices"][0]["message"]["content"]

            # Extract JSON array
            json_start = content.find("[")
            json_end = content.rfind("]") + 1
            if json_start >= 0 and json_end > json_start:
                analyses = json.loads(content[json_start:json_end])
                for a in analyses:
                    idx = a.get("idx", -1)
                    if 0 <= idx < len(batch):
                        batch[idx]["impact"] = a.get("impact", "中性")
                        batch[idx]["direction"] = a.get("direction", "neutral")
                        batch[idx]["sectors"] = ", ".join(a.get("sectors", []))
                        batch[idx]["stocks"] = ", ".join(a.get("stocks", []))
                        batch[idx]["reasoning"] = a.get("reasoning", "")
            else:
                print(f"    Batch {batch_start}: LLM returned non-JSON: {content[:100]}...")
        except Exception as e:
            print(f"    Batch {batch_start} error: {e}")

        analyzed.extend(batch)
        print(f"    [{batch_start+1}-{min(batch_start+batch_size, len(items))}] analyzed")
        if batch_start + batch_size < len(items):
            time.sleep(1)

    return analyzed

def build_douyin_images(items, date_str, out_dir):
    """Generate Douyin-ready portrait images (1080x1920) for carousel posting.
    Returns list of image paths."""
    from PIL import Image, ImageDraw, ImageFont
    import textwrap

    W, H = 1080, 1920
    paths = []

    # Try to use CJK font
    font_paths = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    title_font = None
    body_font = None
    small_font = None

    for fp in font_paths:
        if os.path.exists(fp):
            try:
                title_font = ImageFont.truetype(fp, 56)
                body_font = ImageFont.truetype(fp, 32)
                small_font = ImageFont.truetype(fp, 24)
                break
            except:
                pass

    if not title_font:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    # Color scheme (dark finance theme)
    BG = (18, 18, 36)         # dark navy
    ACCENT = (233, 69, 96)    # red accent
    GOLD = (255, 200, 60)     # gold for highlights
    WHITE = (255, 255, 255)
    LIGHT_GRAY = (180, 180, 190)
    CARD_BG = (30, 30, 55)    # slightly lighter card

    # ---- Slide 0: Title/Cover ----
    img = Image.new("RGB", (W, H), BG)
    draw = ImageDraw.Draw(img)

    # Top accent bar
    draw.rectangle([(0, 0), (W, 8)], fill=ACCENT)

    # Date
    draw.text((80, 200), date_str, fill=LIGHT_GRAY, font=small_font)

    # Title
    title = "每日财经要闻"
    bbox = draw.textbbox((0, 0), title, font=title_font)
    tw = bbox[2] - bbox[0]
    draw.text(((W - tw)//2, 280), title, fill=WHITE, font=title_font)

    # Subtitle
    subtitle = f"共 {len(items)} 条 · {len(set(i['source_id'] for i in items))} 个来源"
    bbox = draw.textbbox((0, 0), subtitle, font=body_font)
    sw = bbox[2] - bbox[0]
    draw.text(((W - sw)//2, 380), subtitle, fill=GOLD, font=body_font)

    # Decorative line
    draw.line([(340, 460), (740, 460)], fill=ACCENT, width=3)

    # Category preview
    by_cat = {}
    for item in items:
        by_cat.setdefault(item["category"], []).append(item)
    cat_order = ["breaking_news", "cn_government", "intl_government",
                 "comprehensive", "global_news", "global_markets",
                 "investment_research", "domestic_research"]
    y = 540
    for cat in cat_order:
        if cat in by_cat:
            cnt = len(by_cat[cat])
            label = CAT_LABELS.get(cat, cat)
            draw.text((120, y), f"{label}　{cnt} 条", fill=LIGHT_GRAY, font=small_font)
            y += 50

    # Footer
    draw.text((80, H - 120), "扫码关注获取每日财经要闻", fill=LIGHT_GRAY, font=small_font)
    draw.text((80, H - 80), "Financial News Collector · 自动生成", fill=(100, 100, 120), font=small_font)

    cover_path = os.path.join(out_dir, f"douyin_00_cover.png")
    img.save(cover_path, quality=95)
    paths.append(cover_path)
    print(f"    Slide 0: Cover")

    # ---- Slides 1-N: Category cards ----
    slide_idx = 1
    for cat in cat_order:
        if cat not in by_cat:
            continue
        cat_items = by_cat[cat][:5]
        label = CAT_LABELS.get(cat, cat)

        img = Image.new("RGB", (W, H), BG)
        draw = ImageDraw.Draw(img)

        # Header
        draw.rectangle([(0, 0), (W, 160)], fill=CARD_BG)
        draw.text((80, 40), label, fill=ACCENT, font=title_font)
        cnt_text = f"{len(cat_items)} 条重要新闻"
        draw.text((80, 110), cnt_text, fill=LIGHT_GRAY, font=small_font)

        y = 220
        for item in cat_items:
            # Card background
            card_h = 300
            if y + card_h > H - 100:
                card_h = H - 100 - y
            draw.rectangle([(40, y), (W-40, y+card_h)], fill=CARD_BG, outline=(60, 60, 80))

            # Source badge
            source = item["source_name"]
            draw.text((80, y + 20), source, fill=GOLD, font=small_font)

            # Impact tag
            impact = item.get("impact", "")
            if impact:
                imp_color = (34, 197, 94) if "利好" in impact else ((239, 68, 68) if "利空" in impact else (180, 180, 190))
                draw.text((80, y + 20), f"【{impact}】", fill=imp_color, font=small_font)
                y_off = 50
            else:
                y_off = 0
            # Title (wrap text)
            title_text = item["title"]
            wrapped = textwrap.fill(title_text, width=30)
            draw.text((80, y + 60 + y_off), wrapped, fill=WHITE, font=body_font)

            # Summary
            summary = item.get("summary", "")[:120]
            if summary:
                wrapped_s = textwrap.fill(summary, width=40)
                draw.text((80, y + 160), wrapped_s, fill=LIGHT_GRAY, font=small_font)

            y += card_h + 30

        # Footer
        draw.text((80, H - 60), f"{slide_idx}  ·  Financial News Collector", fill=(100, 100, 120), font=small_font)

        img_path = os.path.join(out_dir, f"douyin_{slide_idx:02d}_{cat}.png")
        img.save(img_path, quality=95)
        paths.append(img_path)
        print(f"    Slide {slide_idx}: {label}")
        slide_idx += 1

    return paths

def build_docx(items, date_str, docx_path):
    """Generate a formatted DOCX report."""
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Arial'

    # Title
    title = doc.add_heading(f'每日财经新闻聚合报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f'日期: {date_str}　|　新闻: {len(items)} 条　|　自动生成')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    by_cat = {}
    for item in items:
        by_cat.setdefault(item["category"], []).append(item)
    cat_order = ["breaking_news", "cn_government", "intl_government", "comprehensive",
                 "global_news", "global_markets", "investment_research", "domestic_research"]

    for cat in cat_order:
        if cat not in by_cat:
            continue
        doc.add_heading(CAT_LABELS.get(cat, cat), level=1)
        for item in by_cat[cat][:6]:
            title_text = item["title"]
            url = item.get("url", "")
            source = item["source_name"]
            summary = item.get("summary", "")

            # Title as bold hyperlink
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f'• {title_text}')
            run.bold = True
            run.font.size = Pt(9)
            if url:
                run.font.color.rgb = RGBColor(0, 102, 204)

            # Impact tag
            impact = item.get("impact", "")
            if impact:
                p_imp = doc.add_paragraph(f'  【{impact}】')
                p_imp.runs[0].font.size = Pt(9)
                p_imp.runs[0].font.color.rgb = RGBColor(34, 197, 94) if "利好" in impact else (RGBColor(239, 68, 68) if "利空" in impact else RGBColor(136, 136, 136))
            # Source
            reasoning = item.get("reasoning", "")
            src_text = f'  {source}'
            if reasoning:
                src_text += f' | 💡 {reasoning}'
            p2 = doc.add_paragraph(src_text)
            p2.paragraph_format.space_after = Pt(2)
            p2.runs[0].font.size = Pt(7)
            p2.runs[0].font.color.rgb = RGBColor(160, 160, 160)

            # Summary
            if summary and summary != title_text:
                p3 = doc.add_paragraph(f'  {summary[:200]}')
                p3.paragraph_format.space_after = Pt(6)
                p3.runs[0].font.size = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph('本报告由 Financial News Collector 自动生成 · GitHub Actions · 仅供参考')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(160, 160, 160)

    doc.save(docx_path)

def send_email(sender, password, to, subject, body, attachments=None):
    domain = sender.split("@")[-1].lower()
    host, port, ssl = SMTP_CFG.get(domain, (f"smtp.{domain}", 465, True))

    msg = MIMEMultipart()
    msg["From"] = sender
    msg["To"] = to
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain", "utf-8"))

    if attachments:
        for att_path in attachments:
            if att_path and os.path.exists(att_path):
                with open(att_path, "rb") as f:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition",
                                f'attachment; filename="{os.path.basename(att_path)}"')
                msg.attach(part)

    if ssl:
        server = smtplib.SMTP_SSL(host, port, timeout=30)
    else:
        server = smtplib.SMTP(host, port, timeout=30)
        server.starttls()
    try:
        server.login(sender, password)
        server.sendmail(sender, [to], msg.as_string())
        print(f"  Email sent -> {to}")
    finally:
        server.quit()


def main():
    print(f"[{now_bj().strftime('%Y-%m-%d %H:%M:%S')}] Starting...")
    yesterday = (now_bj() - timedelta(days=1)).strftime("%Y-%m-%d")
    today_str = now_bj().strftime("%Y-%m-%d")

    print(f"\n[1/4] Searching {len(SOURCES)} sources (DDG HTML)...")
    items = search_all(since_date=yesterday, max_per=5)
    hits = len(set(i["source_id"] for i in items))
    print(f"  -> {len(items)} items from {hits}/{len(SOURCES)} sources")

    print(f"\n[2/5] Analyzing capital market impact (SiliconFlow LLM)...")
    items = analyze_impact(items)
    analyzed_count = sum(1 for i in items if i.get("impact"))

    print(f"\n[3/5] Building text report...")
    report = text_report(items, today_str)

    print(f"\n[4/5] Generating PDF...")
    pdf_path = os.path.join(tempfile.gettempdir(), f"每日财经新闻_{today_str}.pdf")
    build_pdf(items, today_str, pdf_path)
    print(f"  PDF: {pdf_path} ({os.path.getsize(pdf_path)/1024:.1f} KB)")

    # Generate DOCX
    docx_path = os.path.join(tempfile.gettempdir(), f"每日财经新闻_{today_str}.docx")
    build_docx(items, today_str, docx_path)
    print(f"  DOCX: {docx_path} ({os.path.getsize(docx_path)/1024:.1f} KB)")

    # Generate Douyin images
    douyin_dir = tempfile.mkdtemp(prefix="douyin_")
    douyin_paths = build_douyin_images(items, today_str, douyin_dir)
    print(f"  Douyin: {len(douyin_paths)} slides generated")

    print(f"\n[5/5] Sending email...")
    sender = os.environ.get("EMAIL_SENDER", "")
    password = os.environ.get("EMAIL_PASSWORD", "")
    recipient = os.environ.get("EMAIL_TO", sender)

    if sender and password:
        all_attachments = [pdf_path, docx_path] + douyin_paths
        send_email(sender, password, recipient,
                   f"📰 每日财经要闻 {today_str}", report, all_attachments)
    else:
        print("  SKIP: no email config")
        print(report[:500])

    print(f"\n[{now_bj().strftime('%Y-%m-%d %H:%M:%S')}] Done!")


if __name__ == "__main__":
    main()
